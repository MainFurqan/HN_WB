"""CV / resume parser.

Accepts a PDF, DOCX, or plain-text CV upload, extracts text, and asks the LLM
to map it onto the 6 profile sections used by the youth flow.

The output shape matches the Profile type on the frontend:
  { about, education, work, selfTaught, tools, aspirations }
"""
from __future__ import annotations
import io
from typing import Any
from ..llm import chat_json
from ..config import settings


CV_SYSTEM_PROMPT = """You are a careful CV/resume reader. You will be given the \
full text of a young person's CV. Your job is to extract their actual content \
and place it into a structured profile with 6 sections, then return JSON.

Sections and what to put in each:
- about: 1-3 sentences. Name, age (if stated), city/country, languages spoken. \
Pull from the header / personal-details / contact section.
- education: degrees, schools, field of study, dates, GPA. One section per \
degree, separated by blank lines if multiple.
- work: job titles, employers, dates, and 1-3 bullet points of what they \
actually did. Keep verbs and impact. Multiple jobs separated by blank lines.
- selfTaught: online courses, certifications, books, side projects, mentors. \
ONLY content that is clearly outside formal education.
- tools: programming languages, frameworks, software, hardware, machinery — \
whatever they list as tools / tech / skills (not soft skills).
- aspirations: career objective or summary line if present. Often the very \
first paragraph of the CV. If absent, leave empty string.

HARD RULES:
- Use the user's own wording where possible. Do not invent or embellish.
- If a section has no matching CV content, return an empty string for it.
- Never fabricate names, dates, or employers.
- Be concise. Do not pad.
- Return ONLY the JSON object — no commentary, no markdown.

Output JSON shape (illustrative — real values come from the CV):
{
  "about": "...",
  "education": "...",
  "work": "...",
  "selfTaught": "...",
  "tools": "...",
  "aspirations": ""
}"""


def extract_text_from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def extract_text_from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if name.endswith(".docx"):
        return extract_text_from_docx(data)
    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", errors="replace")
    # Best effort: try as text
    return data.decode("utf-8", errors="replace")


def parse_cv_to_profile(filename: str, data: bytes) -> dict[str, Any]:
    text = extract_text(filename, data)
    text = text.strip()
    if not text:
        return {
            "about": "",
            "education": "",
            "work": "",
            "selfTaught": "",
            "tools": "",
            "aspirations": "",
            "raw_text_chars": 0,
        }
    # Cap input to keep cost bounded.
    if len(text) > 18000:
        text = text[:18000]

    user_prompt = f"CV CONTENT (verbatim):\n\n{text}"
    result = chat_json(
        messages=[
            {"role": "system", "content": CV_SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        model=settings.openai_model_fast,
        temperature=0.1,
    )

    out = {
        "about": str(result.get("about") or ""),
        "education": str(result.get("education") or ""),
        "work": str(result.get("work") or ""),
        "selfTaught": str(result.get("selfTaught") or result.get("self_taught") or ""),
        "tools": str(result.get("tools") or ""),
        "aspirations": str(result.get("aspirations") or ""),
        "raw_text_chars": len(text),
    }
    return out
